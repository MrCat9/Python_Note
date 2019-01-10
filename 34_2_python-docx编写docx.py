# -*- coding: utf-8 -*-
# https://blog.csdn.net/sinat_30711195/article/details/80725435
# python-docx 文档 https://python-docx.readthedocs.io/en/latest/

from docx import Document
from docx.shared import Inches  # 图片大小
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 居中
from docx.shared import RGBColor  # 颜色
from docx.oxml.ns import qn  # 字体
from docx.shared import Pt  # 字体大小


def gen_doc_report(doc_save_path, info_list):
    if info_list.__len__() == 1:
        info_dict = info_list[0]

        document = Document()

        document.styles['Normal'].font.name = '宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        head = document.add_heading('这里是标题', 0)  # 标题
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中

        # 要写在表格中的内容
        name1 = '姓名：{}'.format(info_dict.get('name1', ''))
        name2 = '姓名：{}'.format(info_dict.get('name2', ''))
        name3 = '姓名：{}'.format(info_dict.get('name3', ''))
        name4 = '姓名：{}'.format(info_dict.get('name4', ''))
        name5 = '姓名：{}'.format(info_dict.get('name5', ''))

        des1 = '简介：{}'.format(info_dict.get('des1', ''))
        des2 = '简介：{}'.format(info_dict.get('des2', ''))

        records = (
            name1, name2, name3,
            name4, name5,

        )

        rows_int = 3  # 行数
        cols_int = 2  # 列数
        table = document.add_table(rows=rows_int, cols=cols_int)  # 生成表格
        item_no = 0
        # 把records中的内容往表格里填
        for i in range(rows_int):
            row = table.rows[i].cells
            for j in range(cols_int):
                paragraphs = row[j].paragraphs[0]
                run = paragraphs.add_run(records[item_no])
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(11)  # 字体大小设置，和word里面的字号相对应
                item_no += 1
                if item_no >= records.__len__():
                    break

        document.add_paragraph('______________________________________________________________________________')

        # 插入图片
        imagePath = info_dict.get('imagePath', '')
        if imagePath:
            if imagePath[-4:] == '.jpg':
                document.add_picture(imagePath, width=Inches(5.5))  # 传图片路径

        # run = document.add_paragraph().add_run('dse1:')
        # run.bold = True  # 加粗
        document.add_paragraph(
                'des1:', style='List Bullet'
            )
        run = document.add_paragraph().add_run(des1)

        # run = document.add_paragraph().add_run('des2:')
        # run.bold = True
        document.add_paragraph(
            'des2:', style='List Bullet'
        )
        run = document.add_paragraph().add_run(des2)

        document.add_paragraph('______________________________________________________________________________')

        time1 = 'time1:{}'.format("2018/2/26  09:35:13")
        time2 = 'time2{}'.format("2018/2/26  11:15:03")

        records = (
            time1, time2,

        )

        rows_int = 1  # 行数
        cols_int = 2  # 列数
        table = document.add_table(rows=rows_int, cols=cols_int)
        item_no = 0
        for i in range(rows_int):
            row = table.rows[i].cells
            for j in range(cols_int):
                paragraphs = row[j].paragraphs[0]
                run = paragraphs.add_run(records[item_no])
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(11)  # 字体大小设置，和word里面的字号相对应
                item_no += 1
                if item_no >= records.__len__():
                    break

        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        run = p.add_run('仅供参考')
        run.bold = True

        # document.add_page_break()  # 分页

        document.save(doc_save_path)

        return 1, "success"

    else:
        return 0, "fail"


if __name__ == '__main__':
    list1 = [{
        'name1': 'name1', 'name2': 'name2', 'name3': 'name3', 'name4': 'name4',
        'name5': 'name5', 'name6': 'name6',
        'imagePath': 'F:\\pycharm\\···\\test_jpg.jpg',
        'des1': 'des1des1des1des1des1des1des1des1des1des1des1des1',
        'des2': 'des2des2des2des2des2des2des2des2',

    }]

    status, msg = gen_doc_report(doc_save_path='F:\\pycharm\\···\\demo.docx', info_list=list1)
    print(status)
    print(msg)
